# -*- coding: utf-8 -*-
"""Generate Analiza_kazusow.docx using python-docx."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUTPUT_PATH = r"C:\Users\Dominik\Desktop\Kalkulator prawniczy\cywil\v.2.1.6\Analiza_kazusow.docx"

doc = Document()

# ---- Page setup: A4, 2.5 cm margins ----
section = doc.sections[0]
section.page_width = 11906  # DXA (twips) for A4
section.page_height = 16838
margin = Cm(2.5)
section.top_margin = margin
section.bottom_margin = margin
section.left_margin = margin
section.right_margin = margin

FONT_NAME = "Calibri"

def set_run_font(run, size_pt, bold=False, italic=False, font_name=FONT_NAME):
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size_pt)
    run.font.name = font_name
    # Ensure font applies to complex scripts too
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)

def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:after'), '240')
    pPr.append(spacing)
    run = p.add_run(text)
    set_run_font(run, 16, bold=True)
    return p

def add_intro(doc, text):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:after'), '240')
    pPr.append(spacing)
    run = p.add_run(text)
    set_run_font(run, 11, italic=True)
    return p

def add_kazus_heading(doc, text):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '240')
    spacing.set(qn('w:after'), '120')
    pPr.append(spacing)
    run = p.add_run(text)
    set_run_font(run, 13, bold=True)
    return p

def add_subsection_heading(doc, text):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '120')
    spacing.set(qn('w:after'), '60')
    pPr.append(spacing)
    run = p.add_run(text)
    set_run_font(run, 11, bold=True, italic=True)
    return p

def add_numbered_item(doc, number, text):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:after'), '60')
    pPr.append(spacing)
    run = p.add_run(f"{number}. {text}")
    set_run_font(run, 11)
    return p

# ===== TITLE =====
add_title(doc, "Zastrzeżenia merytoryczne i językowe – analiza kazusów z prawa pracy i ubezpieczeń społecznych")

# ===== INTRO =====
add_intro(doc, "Niniejszy dokument zawiera zastrzeżenia merytoryczne oraz językowe do kazusów i odpowiedzi modelowych z zakresu prawa pracy i ubezpieczeń społecznych. Zastrzeżenia mają charakter sygnalizacyjny – nie wprowadzono do tekstu żadnych samodzielnych poprawek.")

# ===== KAZUS 1 =====
add_kazus_heading(doc, "KAZUS 1. Roszczenia uzupełniające w przypadku rozwiązania umowy o pracę bez wypowiedzenia z winy pracownika")

add_subsection_heading(doc, "Zastrzeżenia językowe:")
add_numbered_item(doc, 1, 'W zdaniu „Zaistnienie odpowiedzialności pracodawcy na podstawie art. 415 k.c. jest możliwa jedynie w przypadku…” – użyto formy „jest możliwa", która jest niezgodna z zasadami gramatyki języka polskiego. Podmiotem zdania jest rzeczownik nijaki „zaistnienie", który wymaga orzecznika w formie nijakiej: „jest możliwe".')

add_subsection_heading(doc, "Zastrzeżenia merytoryczne:")
add_numbered_item(doc, 1, "Odpowiedź prawidłowo identyfikuje konieczność wykazania winy umyślnej pracodawcy w postaci zamiaru bezpośredniego kierunkowego, jednak nie rozwinięto zagadnienia, w jaki sposób fakt politycznej motywacji zwolnienia (zmiana władzy w gminie po wyborach) przekłada się na wykazanie tego zamiaru w kontekście dowodowym. Warto było wskazać, że motywacja polityczna może stanowić okoliczność ułatwiającą dowód zamiaru kierunkowego, ale sama w sobie go nie przesądza.")
add_numbered_item(doc, 2, "Odpowiedź nie odnosi się wprost do kwestii związku przyczynowego pomiędzy niezgodnym z prawem zwolnieniem a konkretnie wskazywaną przez pracownika szkodą – w szczególności do rocznej niezdolności do poszukiwania pracy spowodowanej depresją. Związek przyczynowy między rozwiązaniem umowy a rozstrojem zdrowia psychicznego jest zagadnieniem wymagającym dowodu z opinii biegłego, co stanowi istotną trudność procesową dla powoda; odpowiedź tej kwestii nie sygnalizuje.")
add_numbered_item(doc, 3, "Nie wskazano wyraźnie, że art. 58 k.p. ogranicza odszkodowanie wyłącznie do trzech miesięcznych wynagrodzeń i że właśnie ta dysproporcja względem faktycznej szkody stanowi ratio legis dopuszczenia odpowiedzialności uzupełniającej – co jest istotne dla oceny stanowisk stron.")

# ===== KAZUS 2 =====
add_kazus_heading(doc, "KAZUS 2. Wynagrodzenie za pracę realne a wskazane na umowie o pracę")

add_subsection_heading(doc, "Zastrzeżenia językowe:")
add_numbered_item(doc, 1, 'W zdaniu „Porozumienie te, wbrew przedstawionemu w kazusie argumentowi, może zostać zawarte również w formie innej niż pisemna" – użyto formy „Porozumienie te", która jest błędna. Rzeczownik „porozumienie" jest rodzaju nijakiego i wymaga zaimka wskazującego „to": „Porozumienie to".')

add_subsection_heading(doc, "Zastrzeżenia merytoryczne:")
add_numbered_item(doc, 1, 'W treści kazusu pracownica nazwana jest najpierw „Genowefa W." (opis stanu faktycznego), a następnie „Katarzyna W." (przy opisie złożenia pozwu). Odpowiedź modelowa konsekwentnie operuje imieniem „Katarzyna W.", pomijając tę niespójność bez komentarza. Przy analizie stanu faktycznego niespójność nazewnictwa powinna zostać odnotowana – jej przeoczenie może prowadzić do wątpliwości, czy kazus dotyczy jednej, czy dwóch różnych osób.')
add_numbered_item(doc, 2, "Odpowiedź słusznie wskazuje na możliwość zastosowania art. 8 k.p., jednak nie wyjaśnia dostatecznie, dlaczego klauzula generalna nadużycia prawa podmiotowego mogłaby uzasadniać oddalenie powództwa – a nie jedynie jego miarkowanie lub modyfikację zakresu zasądzenia. Warto było zaznaczyć, że art. 8 k.p. nie prowadzi do oddalenia powództwa co do zasady, lecz może jedynie ograniczyć zakres uwzględnionego żądania.")
add_numbered_item(doc, 3, "Odpowiedź nie rozważa, czy możliwe jest powołanie się na nieważność umowy (lub jej części) ze względu na sprzeczność z zasadami współżycia społecznego (art. 58 § 2 k.c. w zw. z art. 300 k.p.) w zakresie, w jakim jej celem było obejście przepisów prawa podatkowego i ubezpieczeniowego. Jest to zagadnienie istotne dla rozstrzygnięcia sprawy i wymagało co najmniej sygnalizacji.")

# ===== KAZUS 3 =====
add_kazus_heading(doc, "KAZUS 3. Odpowiedzialność za składki przypadające pracownikowi/zleceniobiorcy w przypadku ustalenia tytułu do podlegania ubezpieczeniom")

add_subsection_heading(doc, "Zastrzeżenia merytoryczne:")
add_numbered_item(doc, 1, 'W treści kazusu wskazano, że ZUS wydał decyzję ustalającą „tytuł do ubezpieczenia (umowa o pracę)" w stosunku do osoby wykonującej umowę o dzieło. Odpowiedź nie zwraca uwagi, że przekwalifikowanie umowy o dzieło bezpośrednio na stosunek pracy jest w praktyce orzeczniczej niezwykle rzadkie i kontrowersyjne. ZUS z reguły kwalifikuje umowy o dzieło jako umowy zlecenia (art. 6 ust. 1 pkt 4 u.s.u.s.), nie zaś jako umowy o pracę. Pominięcie tej kwestii stanowi istotne uproszczenie.')
add_numbered_item(doc, 2, "Odpowiedź nie sygnalizuje, że ewentualna nieświadomość wykonawcy co do zasad podlegania ubezpieczeniom (jako typowego zleceniobiorcy/wykonawcy dzieła) może mieć znaczenie dla oceny roszczenia na podstawie art. 8 k.p. – jako okoliczność mogąca przemawiać za odmową uwzględnienia roszczenia pracodawcy.")
add_numbered_item(doc, 3, "Odpowiedź nie wskazuje, że prawidłową drogą dochodzenia zwrotu niezapłaconych składek przez pracodawcę-płatnika powinno być, co do zasady, potrącenie z bieżących należności pracownika (art. 87 k.p.), a roszczenie z bezpodstawnego wzbogacenia aktualizuje się dopiero wtedy, gdy potrącenie nie jest możliwe – co ma znaczenie dla dopuszczalności powództwa.")

# ===== KAZUS 4 =====
add_kazus_heading(doc, "KAZUS 4. Renta rodzinna")

add_subsection_heading(doc, "Zastrzeżenia merytoryczne:")
add_numbered_item(doc, 1, "Odpowiedź prawidłowo wskazuje, że Barbara K. mogła działać w imieniu pełnoletniego syna wyłącznie jako pełnomocnik, nie zaś przedstawiciel ustawowy, jednak nie precyzuje formy, w jakiej potwierdzenie czynności pełnomocnika przez Damiana powinno nastąpić w postępowaniu administracyjnym (pisemne oświadczenie lub ustne do protokołu – art. 33 k.p.a.).")
add_numbered_item(doc, 2, "Odpowiedź nie odnosi się do tego, jaki skutek procesowy wywoła brak potwierdzenia działania matki przez pełnoletniego syna – w szczególności, czy ZUS powinien wezwać Damiana do potwierdzenia lub do samodzielnego złożenia wniosku w trybie art. 64 § 2 k.p.a., zanim wyda decyzję odmowną w tej części.")
add_numbered_item(doc, 3, "Odpowiedź nie sygnalizuje, że Damian (23 lata) może być uprawniony do renty rodzinnej wyłącznie pod warunkiem nauki, a prawo do renty wygaśnie z chwilą ukończenia przez niego 25 roku życia lub przerwania nauki (art. 68 ust. 1 pkt 2 u.e.r.f.u.s.) – te konsekwencje temporalne warto było wyeksponować.")
add_numbered_item(doc, 4, "Odpowiedź nie omawia ewentualnego prawa Barbary K. do dodatku dla sieroty zupełnej ani do innych dodatkowych świadczeń pochodnych od renty rodzinnej, które mogłyby być istotne z perspektywy kompleksowej oceny uprawnień wnioskodawczyni.")

# ===== KAZUS 5 =====
add_kazus_heading(doc, "KAZUS 5. Renta z tytułu niezdolności do pracy")

add_subsection_heading(doc, "Zastrzeżenia merytoryczne:")
add_numbered_item(doc, 1, 'Odpowiedź stwierdza, że wymagany okres składkowy i nieskładkowy dla Kazimierza D. wynosi „co najmniej 3 lata", co jest prawidłowe w świetle art. 58 ust. 1 pkt 3 u.e.r.f.u.s. (niezdolność do pracy powstała w wieku powyżej 22 do 25 lat). Jednakże odpowiedź nie wskazuje wprost numeru właściwego punktu art. 58 ust. 1, co obniża precyzję wywodu prawniczego.')
add_numbered_item(doc, 2, "Odpowiedź nie omawia możliwości przyznania renty szkoleniowej (art. 60 u.e.r.f.u.s.) w przypadku, gdyby Lekarz Orzecznik ZUS orzekł celowość przekwalifikowania zawodowego Kazimierza D. Biorąc pod uwagę charakter schorzenia (trudności ze stabilnym trzymaniem ciężkiego sprzętu), przekwalifikowanie zawodowe jest realną perspektywą i instytucja renty szkoleniowej powinna zostać omówiona.")
add_numbered_item(doc, 3, "Odpowiedź nie porusza kwestii ewentualnego prawa do świadczenia rehabilitacyjnego jako instrumentu poprzedzającego przyznanie renty, choć w stanie faktycznym kazusu wzmiankuje się, że Kazimierz D. już je pobierał. Brak wzmianki o możliwości ponownego ubiegania się o świadczenie rehabilitacyjne (gdyby renta nie została przyznana) stanowi lukę w analizie dostępnych świadczeń.")

# ===== KAZUS 6 =====
add_kazus_heading(doc, "KAZUS 6. Naruszenie dóbr osobistych pracownika a mobbing")

add_subsection_heading(doc, "Zastrzeżenia merytoryczne:")
add_numbered_item(doc, 1, 'Odpowiedź formułuje tezę, że „katalog dóbr osobistych pracownika nie jest tożsamy z katalogiem znajdującym się w Kodeksie cywilnym". Teza ta jest kontrowersyjna i wymaga ostrożniejszego sformułowania. Dominujące orzecznictwo Sądu Najwyższego przyjmuje, że art. 111 k.p. nie kreuje odrębnego, autonomicznego katalogu dóbr osobistych pracownika, lecz stanowi normę odsyłającą do ogólnej ochrony dóbr osobistych na gruncie art. 23–24 k.c. w zw. z art. 300 k.p. Prezentowanie poglądu o odmienności katalogów bez wyraźnego zaznaczenia jego sporności może wprowadzać czytelnika w błąd.')
add_numbered_item(doc, 2, 'Odpowiedź prawidłowo sygnalizuje obowiązek pouczenia pracownika przez przewodniczącego (art. 477 zd. 2 k.p.c.), jednak nie akcentuje dostatecznie, że pouczenie to ma wyłącznie charakter informacyjny i nie uprawnia sądu do orzekania ponad żądanie (art. 321 § 1 k.p.c.). Brak tej uwagi może sugerować, że sąd może niejako „z urzędu" przekształcić podstawę roszczenia.')

# ===== KAZUS 7 =====
add_kazus_heading(doc, "KAZUS 7. Wynagrodzenie za nadgodziny w przypadku sędziego")

add_subsection_heading(doc, "Zastrzeżenia językowe:")
add_numbered_item(doc, 1, 'W zdaniu „treść wyroku zdaje się wskazywać, że w istocie sędzia nie jest w ogóle zobowiązany do zachowywania norm pracy, bowiem orzecznicy nie są zobowiązania do rozliczania się z pracodawcą" – użyto błędnej formy fleksyjnej „zobowiązania" zamiast poprawnej „zobowiązani". Jest to błąd gramatyczny w zakresie odmiany przymiotnika/imiesłowu przymiotnikowego biernego przez rodzaj.')

add_subsection_heading(doc, "Zastrzeżenia merytoryczne:")
add_numbered_item(doc, 1, "Odpowiedź słusznie powołuje uchwałę SN z 8.04.2009 r. (II PZP 2/09) oraz wyrok TK z 7.05.2013 r. (SK 11/11), jednak nie wskazuje, że od czasu tych orzeczeń w doktrynie formułowane są postulaty de lege ferenda dotyczące wprowadzenia maksymalnego obciążenia sędziów sprawami. Sygnalizacja tej luki legislacyjnej wzbogaciłaby analizę.")
add_numbered_item(doc, 2, "Odpowiedź nie odnosi się do zagadnienia zgodności art. 83 u.s.p. z prawem Unii Europejskiej, w szczególności z Dyrektywą 2003/88/WE o czasie pracy, pomimo że kwestia ta była poruszana w piśmiennictwie. Rozstrzygnięcie TK dotyczyło zgodności z Konstytucją RP, nie z prawem unijnym.")

# ===== KAZUS 8 =====
add_kazus_heading(doc, "KAZUS 8. Zatrudnienie osoby z niepełnosprawnością")

add_subsection_heading(doc, "Zastrzeżenia merytoryczne:")
add_numbered_item(doc, 1, 'Błąd merytoryczny: w odpowiedzi wskazano, że implementacja przepisów antydyskryminacyjnych nastąpiła na podstawie „dyrektywy Rady 2000/78/WE z 27 listopada 2003 roku". Jest to nieprawidłowe – dyrektywa Rady 2000/78/WE pochodzi z dnia 27 listopada 2000 roku, nie 2003 roku. Podanie błędnej daty aktu prawa Unii Europejskiej stanowi istotne uchybienie merytoryczne.')
add_numbered_item(doc, 2, "Odpowiedź nie omawia obowiązku pracodawcy wprowadzenia racjonalnych usprawnień na rzecz pracownika z niepełnosprawnością (art. 23a ustawy o rehabilitacji zawodowej i społecznej oraz zatrudnianiu osób niepełnosprawnych), który to obowiązek stanowi jedno z kluczowych narzędzi ochrony przed dyskryminacją pośrednią osób niepełnosprawnych. Pominięcie tego zagadnienia jest istotną luką w analizie prawnej.")
add_numbered_item(doc, 3, 'Odpowiedź nie wskazuje na możliwość uzyskania przez pracodawcę dofinansowania do wynagrodzenia pracownika z niepełnosprawnością ze środków Państwowego Funduszu Rehabilitacji Osób Niepełnosprawnych (art. 26a i n. ustawy o rehabilitacji). Ta okoliczność jest bezpośrednio istotna dla pytania Szymona M. o to, czy zatrudnienie Rafała S. „wpłynie na harmonogram pracy w zakładzie" – a w istocie może wpłynąć na nie korzystnie (finansowo).')
add_numbered_item(doc, 4, "Odpowiedź nie sygnalizuje, że zrzeczenie się przez pracownika przysługujących mu ustawowych uprawnień (np. skróconego czasu pracy, dodatkowego urlopu) jest co do zasady niedopuszczalne bez spełnienia wymogów formalnych przewidzianych w ustawie o rehabilitacji (w szczególności wymogu zgody lekarskiej), co jest istotne z punktu widzenia pytania o dopuszczalność zrzeczenia się tych przywilejów.")

doc.save(OUTPUT_PATH)
print(f"Saved: {OUTPUT_PATH}")
